from saga.base_file.File import File
from docx import Document
from os import remove

def parse_word_file(file_id, file_name, file_path):
    document = Document(file_path)
    # for now, we model a word document as a list of paragraphs
    file_contents = [p.text for p in document.paragraphs]
    return File(file_id, "word", file_path, file_name, file_contents)

def write_word_file(file):
    document = Document()

    for p_text in file.file_contents.mixed_data_type.mixed_data_type:
        p = document.add_paragraph(p_text)

    remove(file.file_path)
    document.save(file.file_path)